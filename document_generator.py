from docx import Document
from docx.shared import Pt
import tempfile
import os
import logging
import re
from question_parser import identify_questions

logger = logging.getLogger(__name__)

def apply_custom_formatting(doc):
    """
    Apply custom formatting to the document:
      1. Text between ** and ** is made bold. 
         If the inner text starts with "Question", set its font size to 14 pt.
      2. If a paragraph starts with an asterisk (*) followed by whitespace,
         remove the marker and set the paragraph style to "List Bullet".
    """
    # Regular expression for bold marker: captures text between ** and **
    bold_pattern = re.compile(r'(\*\*.*?\*\*)')
    
    for para in doc.paragraphs:
        original_text = para.text
        
        # Check if the paragraph starts with a bullet marker (e.g., "* ")
        bullet_match = re.match(r'^\*\s+', original_text)
        is_bullet = bool(bullet_match)
        if is_bullet:
            # Remove the bullet marker from the text
            text = re.sub(r'^\*\s+', '', original_text)
        else:
            text = original_text
        
        # Remove all runs from the paragraph so we can rebuild it cleanly.
        p_element = para._element
        for run in list(para.runs):
            p_element.remove(run._element)
        
        # Split the text while preserving the bold markers.
        parts = re.split(bold_pattern, text)
        
        # Rebuild the paragraph by adding new runs with formatting.
        for part in parts:
            # Check if the part matches the **...** pattern.
            m = re.match(r'\*\*(.*?)\*\*', part)
            if m:
                inner_text = m.group(1)
                run = para.add_run(inner_text)
                run.bold = True
                # If the bold text starts with "Question", enlarge the font size.
                if inner_text.strip().startswith("Question"):
                    run.font.size = Pt(14)
            else:
                para.add_run(part)
        
        # If this paragraph originally started with a bullet marker, update its style.
        if is_bullet:
            para.style = 'List Bullet'


def extract_questions(text):
    pattern = r"(\n\n\*\*Question.*?)(?=\n\n\*\*Question|$)"
    matches = re.findall(pattern, text, re.DOTALL)
    return [match.strip() for match in matches]


def create_document(subject, questions_text, answers):
    """
    Create a formatted Word document with questions and answers.
    The document is further processed to apply custom formatting:
      1. Bold text formatting for text between ** and **.
         (If the bold text begins with "Question", the font is enlarged.)
      2. Convert paragraphs starting with an asterisk (*) into bullet points.
    """
    try:
        doc = Document()

        # Add title
        title = doc.add_heading(f'{subject} - Exam Answers', 0)
        title.alignment = 1  # Center alignment

        #----------------New Format-------#

        cleaned_text = '\n\n**Question 1' + answers.split('\n\n**Question 1', 1)[-1]

        answer_sections = extract_questions(cleaned_text)

        print("Answer sections:", answer_sections)

        for ans in answer_sections:
            doc.add_paragraph(ans)
            doc.add_paragraph()
            doc.add_paragraph()
        #-----------------END-------------#

        # Apply custom formatting to process patterns in all paragraphs.
        apply_custom_formatting(doc)

        # Save the document to a temporary path.
        temp_path = os.path.join(tempfile.gettempdir(), 'exam_answers.docx')
        doc.save(temp_path)

        return temp_path
    except Exception as e:
        logger.error(f"Error creating document: {str(e)}")
        raise Exception("Failed to create document")


